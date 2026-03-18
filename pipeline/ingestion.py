import fitz  # PyMuPDF
from docx import Document
import os


# ================= PDF EXTRACTION =================

def extract_pdf_structured(pdf_path):
    """
    Extract PDF as structured paragraphs (no OCR).
    """
    content = []
    doc = fitz.open(pdf_path)

    for page in doc:
        text = page.get_text()

        if text:
            lines = text.split("\n")

            for line in lines:
                line = line.strip()
                if line:
                    content.append({
                        "type": "paragraph",
                        "text": line
                    })

    return content


# ================= DOCX EXTRACTION =================

def extract_docx_structured(docx_path):
    """
    Extract DOCX preserving paragraphs + tables.
    """
    doc = Document(docx_path)
    content = []

    # -------- PARAGRAPHS --------
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            content.append({
                "type": "paragraph",
                "text": text
            })

    # -------- TABLES --------
    for table in doc.tables:
        rows_data = []

        for row in table.rows:
            row_data = []
            for cell in row.cells:
                row_data.append(cell.text.strip())
            rows_data.append(row_data)

        content.append({
            "type": "table",
            "rows": rows_data
        })

    return content


# ================= MAIN INGESTION =================

def extract_content(file_path):
    """
    Main ingestion function used by pipeline.
    Returns structured content (NOT raw text).
    """

    file_name = os.path.basename(file_path)
    ext = file_name.lower().split(".")[-1]

    if ext == "pdf":

        content = extract_pdf_structured(file_path)

        method = "pdf_structured"

    elif ext == "docx":

        content = extract_docx_structured(file_path)

        method = "docx_structured"

    else:
        raise ValueError("Unsupported file type")

    return {
        "file_name": file_name,
        "file_type": ext,
        "content": content,   # 🔥 IMPORTANT: structured, not raw_text
        "extraction_method": method
    }
